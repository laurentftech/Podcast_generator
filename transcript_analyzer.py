"""
Module for analyzing transcripts using Gemini API and generating DOCX documents.
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.genai as genai
from google.genai import types
from generate_podcast import setup_logging
from utils import get_asset_path

logger = setup_logging()


def get_analysis_prompt_path():
    """
    Returns the path to the analysis prompt configuration file.
    Checks in order:
    1. ./config/analysis_prompt.txt (for Docker and local config directory)
    2. app data directory (user-editable location)
    3. asset path (for bundled apps)
    """
    from utils import get_app_data_dir

    # First check in config directory (Docker and local development)
    config_dir_prompt = os.path.join(os.getcwd(), "config", "analysis_prompt.txt")
    if os.path.exists(config_dir_prompt):
        return config_dir_prompt

    # Then check in app data directory (user-editable location)
    app_data_prompt = os.path.join(get_app_data_dir(), "analysis_prompt.txt")
    if os.path.exists(app_data_prompt):
        return app_data_prompt

    # Finally check in asset path (bundled with app)
    asset_prompt = get_asset_path("analysis_prompt.txt")
    if asset_prompt and os.path.exists(asset_prompt):
        return asset_prompt

    return None


def load_analysis_prompt():
    """
    Loads the analysis prompt from the configuration file.
    Returns None if the file doesn't exist.
    """
    prompt_path = get_analysis_prompt_path()
    if not prompt_path:
        return None

    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"Error reading analysis prompt file: {e}")
        return None


def generate_prompt(text: str) -> str:
    """
    Génère le prompt pour l'API Gemini.
    Loads the prompt template from the configuration file and appends the transcript.

    Raises:
        ValueError: If the prompt configuration file is not found or cannot be read.
    """
    prompt_template = load_analysis_prompt()
    if not prompt_template:
        raise ValueError("Analysis prompt configuration file not found. Please create 'analysis_prompt.txt' in the application directory.")

    return f"{prompt_template}\n\nTranscript:\n{text}"


def analyze_transcript_with_gemini(transcript: str, api_key: str = None) -> str:
    """
    Analyzes a transcript using the Gemini API.

    Args:
        transcript: The transcript text to analyze
        api_key: Gemini API key (will use environment variable if not provided)

    Returns:
        The analysis response from Gemini

    Raises:
        ValueError: If no API key is found
        Exception: If the API call fails
    """
    if not api_key:
        api_key = os.environ.get("GEMINI_API_KEY")

    if not api_key:
        raise ValueError("GEMINI_API_KEY not found in environment variables")

    # Get the model name from environment variable or use default
    model_name = os.environ.get("GEMINI_ANALYSIS_MODEL", "gemini-2.5-flash")

    try:
        client = genai.Client(api_key=api_key)
        prompt = generate_prompt(transcript)

        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.7,
                max_output_tokens=2048,
            )
        )

        if response.text:
            return response.text
        else:
            raise Exception("Empty response from Gemini API")

    except Exception as e:
        logger.error(f"Error analyzing transcript with Gemini: {e}", exc_info=True)
        raise


def create_docx_from_analysis(analysis_text: str, output_path: str):
    """
    Creates a well-formatted DOCX document from the Gemini analysis.

    Args:
        analysis_text: The analysis text from Gemini (with markdown-style formatting)
        output_path: Path where to save the DOCX file
    """
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Process the text line by line
    lines = analysis_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if the line contains bold text (**text**)
        bold_pattern = r'\*\*(.+?)\*\*'

        if '**' in line:
            # Create a paragraph
            p = doc.add_paragraph()

            # Split the line by bold markers
            parts = re.split(bold_pattern, line)

            is_bold = False
            for i, part in enumerate(parts):
                if not part:
                    continue

                # Alternate between normal and bold
                if i % 2 == 1:  # Odd indices are bold
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(12)
                else:  # Even indices are normal
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.style = 'Normal'

    # Save the document
    doc.save(output_path)
    logger.info(f"DOCX document saved to: {output_path}")


def generate_analysis_docx(transcript: str, output_path: str = None, api_key: str = None) -> str:
    """
    Complete workflow: analyze transcript with Gemini and create a DOCX document.

    Args:
        transcript: The transcript text to analyze
        output_path: Path where to save the DOCX file (auto-generated if not provided)
        api_key: Gemini API key (will use environment variable if not provided)

    Returns:
        Path to the generated DOCX file
    """
    # Analyze the transcript
    analysis = analyze_transcript_with_gemini(transcript, api_key)

    # Generate output path if not provided
    if not output_path:
        import tempfile
        output_path = os.path.join(tempfile.gettempdir(), f"script_analysis_{os.urandom(4).hex()}.docx")

    # Create the DOCX
    create_docx_from_analysis(analysis, output_path)

    return output_path
